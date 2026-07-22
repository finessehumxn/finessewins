"""
FinesseWins — RFP Shredder

The build that moves FinesseWins from "AI writing tool" to "proposal system a
contracting professional respects." Three passes:

  1. SHRED      — read the actual solicitation, extract every requirement,
                  submission instruction (Section L), and evaluation criterion
                  (Section M) into a structured requirements graph.
  2. ANALYZE    — score the user's own documents against each requirement
                  (addressed / partial / missing) → a compliance matrix + a
                  single coverage %. This is the "watch 60% become 95%" moment.
  3. STRENGTHEN — rewrite a weak/missing section grounded in BOTH the
                  requirement text and the user's real content, mirroring the
                  solicitation's vocabulary so it reads as responsive.

Two hard rules that keep it honest and defensible:
  • STRENGTHEN never invents past performance, certifications, staff, or
    capabilities the user did not provide — a hallucinated qualification in a
    federal proposal is a False Claims Act problem, not a UX nit. Missing facts
    are surfaced as gaps for the user to fill, never fabricated.
  • Every requirement and every suggestion traces back to its source text.
"""
from __future__ import annotations

import json
import os
import re
from typing import List, Optional

import io

from langchain_anthropic import ChatAnthropic
from langchain_core.messages import HumanMessage

ANTHROPIC_MODEL = os.environ.get("ANTHROPIC_MODEL", "claude-sonnet-5")
_llm = None


class ShredderUnavailable(RuntimeError):
    """Raised when the AI key isn't configured."""


def _get_llm(max_tokens: int = 8000):
    global _llm
    if _llm is None or _llm.max_tokens != max_tokens:
        key = os.environ.get("ANTHROPIC_API_KEY")
        if not key:
            raise ShredderUnavailable(
                "ANTHROPIC_API_KEY is not set — the RFP Shredder needs it to read "
                "and analyze documents. Add it to the backend environment."
            )
        _llm = ChatAnthropic(model=ANTHROPIC_MODEL, anthropic_api_key=key, max_tokens=max_tokens)
    return _llm


def _text_of(resp) -> str:
    """Coerce an LLM response to plain text. Newer Claude models return
    `content` as a list of content blocks rather than a string."""
    c = getattr(resp, "content", resp)
    if isinstance(c, str):
        return c
    if isinstance(c, list):
        parts = []
        for b in c:
            if isinstance(b, str):
                parts.append(b)
            elif isinstance(b, dict):
                parts.append(b.get("text") or b.get("content") or "")
            else:
                parts.append(getattr(b, "text", "") or "")
        return "".join(parts)
    return str(c or "")


def _extract_json(text: str):
    """Parse a JSON object/array from an LLM response, tolerating code fences."""
    if not text:
        return None
    t = text.strip()
    t = re.sub(r"^```(?:json)?\s*|\s*```$", "", t, flags=re.IGNORECASE | re.MULTILINE).strip()
    try:
        return json.loads(t)
    except Exception:
        pass
    # Fall back to the first {...} or [...] block (prefer object — our top level).
    for pat in (r"\{.*\}", r"\[.*\]"):
        m = re.search(pat, t, re.DOTALL)
        if m:
            try:
                return json.loads(m.group())
            except Exception:
                continue
    return None


def _clip(text: str, limit: int) -> str:
    return text if len(text) <= limit else text[:limit] + "\n…[truncated]"


# ── PASS 1: SHRED ────────────────────────────────────────────────

SHRED_SYSTEM = (
    "You are a senior government proposal manager doing a compliance shred of a "
    "solicitation. You are precise, literal, and never invent requirements that "
    "aren't in the document. You know that Section L = instructions to offerors "
    "(what/how to submit, format, page limits), Section M = evaluation criteria "
    "(how they score), and the SOW/PWS/Scope describes the work itself. State and "
    "local RFPs use different labels (Instructions, Submittal Requirements, "
    "Evaluation, Scope of Services) — map them to the same idea."
)


async def shred(rfp_text: str, title: str = "", agency: str = "") -> dict:
    """Extract a structured requirements graph from the raw solicitation text."""
    prompt = f"""Shred this solicitation into a structured compliance checklist.

Return ONLY valid JSON with this exact shape:
{{
  "summary": "2-3 plain-English sentences: what they want and what wins",
  "key_dates": [{{"label": "Questions due", "date": "..."}}, {{"label": "Proposal due", "date": "..."}}],
  "submission": {{"format": "...", "page_limit": "... or null", "copies": "... or null", "how_to_submit": "..."}},
  "requirements": [
    {{
      "id": "R1",
      "text": "the requirement stated crisply and actionably",
      "source": "short verbatim quote or section reference it came from",
      "section": "L | M | SOW | submission | eligibility | other",
      "type": "instruction | evaluation | requirement | submission | eligibility",
      "category": "short tag e.g. 'Technical Approach', 'Past Performance', 'Pricing', 'Certifications'",
      "mandatory": true
    }}
  ]
}}

Rules:
- Capture EVERY distinct requirement, submittal instruction, and evaluation factor. Aim for completeness (typically 15-40 items).
- Keep each "text" one clear obligation. Split compound requirements.
- "source" must be grounded in the document — a SHORT quote (<=12 words) or a section number. Do not invent.
- Keep every field terse. Do not pad. The whole JSON must fit well within the token budget.
- Order: eligibility → submission/format (Section L) → evaluation factors (Section M) → scope/technical requirements.

SOLICITATION TITLE: {title or 'Unknown'}
AGENCY: {agency or 'Unknown'}

SOLICITATION TEXT:
{_clip(rfp_text, 38000)}
"""
    resp = await _get_llm(8000).ainvoke([HumanMessage(content=f"{SHRED_SYSTEM}\n\n{prompt}")])
    data = _extract_json(_text_of(resp)) or {}
    if isinstance(data, list):          # model returned just the array
        data = {"requirements": data}
    reqs = data.get("requirements") or []
    # Normalize + guarantee stable ids.
    for i, r in enumerate(reqs, 1):
        r.setdefault("id", f"R{i}")
        r["id"] = str(r["id"])
        for k in ("text", "source", "section", "type", "category"):
            r.setdefault(k, "")
        r["mandatory"] = bool(r.get("mandatory", True))
    data["requirements"] = reqs
    data.setdefault("summary", "")
    data.setdefault("key_dates", [])
    data.setdefault("submission", {})
    return data


# ── PASS 2: ANALYZE (compliance matrix / gap analysis) ───────────

ANALYZE_SYSTEM = (
    "You are a proposal compliance reviewer building a requirements-traceability "
    "matrix. You compare the offeror's existing material against each solicitation "
    "requirement and judge coverage honestly. You cite the evidence you found; if "
    "there is no evidence, you mark it missing rather than assuming."
)


def _norm_status(s, coverage=None):
    """Map whatever word the model used onto our three-state enum."""
    s = str(s or "").strip().lower()
    if s in ("addressed", "met", "full", "fully addressed", "compliant", "complete", "yes", "strong"):
        return "addressed"
    if s in ("partial", "partially addressed", "weak", "partially", "some", "minimal"):
        return "partial"
    if s in ("missing", "not addressed", "none", "no", "absent", "not met", "not evaluated"):
        return "missing"
    # Unknown label → infer from coverage if we have it.
    if isinstance(coverage, (int, float)):
        return "addressed" if coverage >= 80 else "partial" if coverage >= 30 else "missing"
    return "partial"


def _coerce_coverage(v, status=None):
    try:
        return max(0, min(100, int(round(float(v)))))
    except Exception:
        # coverage was a word like "Full"/"None" or absent → derive from status.
        return {"addressed": 90, "partial": 55, "missing": 5}.get(status, 0)


async def _score_batch(batch: List[dict], docs_blob: str) -> List[dict]:
    req_lines = "\n".join(f'- {r["id"]} [{r.get("category","")}] {r["text"]}' for r in batch)
    prompt = f"""Compare the offeror's documents against each requirement and return ONLY valid JSON:
{{"matrix": [{{"id": "R1", "status": "addressed" | "partial" | "missing", "coverage": <integer 0-100>, "doc": "<filename or null>", "evidence": "<=15-word quote or null", "note": "<one short line: strength, or exactly what's missing>"}}]}}

Rules:
- "status" MUST be exactly one of: addressed, partial, missing. "coverage" MUST be an integer 0-100 (not a word).
- addressed = 80-100 (clearly satisfied), partial = 30-79 (touched but weak/generic), missing = 0-29 (not found).
- Judge intent, not keywords. Be honest. Include EVERY requirement id below exactly once. Keep evidence/note short.

REQUIREMENTS:
{req_lines}

OFFEROR'S DOCUMENTS:
{docs_blob}
"""
    resp = await _get_llm(8000).ainvoke([HumanMessage(content=f"{ANALYZE_SYSTEM}\n\n{prompt}")])
    data = _extract_json(_text_of(resp)) or {}
    return data.get("matrix") or []


async def analyze(requirements: List[dict], user_docs: List[dict], batch_size: int = 12) -> dict:
    """Score each requirement against the user's documents → compliance matrix.

    Requirements are scored in batches so the JSON never truncates on long
    solicitations. user_docs: [{"name": str, "text": str}] (already parsed).
    """
    import asyncio
    docs_blob = _clip("\n\n".join(
        f'### DOCUMENT: {d["name"]}\n{_clip(d.get("text",""), 12000)}' for d in user_docs
    ) or "(no documents provided)", 90000)

    batches = [requirements[i:i + batch_size] for i in range(0, len(requirements), batch_size)]
    results = await asyncio.gather(*[_score_batch(b, docs_blob) for b in batches], return_exceptions=True)

    by_id = {}
    for res in results:
        if isinstance(res, Exception):
            continue
        for m in res:
            by_id[str(m.get("id"))] = m

    out = []
    for r in requirements:
        m = by_id.get(r["id"])
        if not m:
            out.append({"id": r["id"], "status": "missing", "coverage": 0,
                        "doc": None, "evidence": None, "note": "Not evaluated."})
            continue
        status = _norm_status(m.get("status"), m.get("coverage"))
        out.append({
            "id": r["id"],
            "status": status,
            "coverage": _coerce_coverage(m.get("coverage"), status),
            "doc": m.get("doc") or None,
            "evidence": m.get("evidence") or None,
            "note": m.get("note") or "",
        })

    counts = {
        "addressed": sum(1 for m in out if m["status"] == "addressed"),
        "partial": sum(1 for m in out if m["status"] == "partial"),
        "missing": sum(1 for m in out if m["status"] == "missing"),
        "total": len(out),
    }
    coverage_pct = round(sum(m["coverage"] for m in out) / len(out)) if out else 0
    return {"matrix": out, "coverage_pct": coverage_pct, "counts": counts}


# ── PASS 3: STRENGTHEN ───────────────────────────────────────────

STRENGTHEN_SYSTEM = (
    "You are an expert proposal writer strengthening an offeror's own material to "
    "respond to a specific solicitation requirement. Absolute rules:\n"
    "1. NEVER invent facts. Do not add past performance, contracts, dollar values, "
    "dates, staff, certifications, or capabilities the offeror did not provide. "
    "Inventing a qualification in a federal proposal is a False Claims Act risk.\n"
    "2. You may reorganize, sharpen, and re-frame what they DID provide, and mirror "
    "the solicitation's vocabulary and evaluation language so it reads as responsive.\n"
    "3. If the requirement needs a fact the offeror hasn't given you, DO NOT make it "
    "up — flag it as a gap the user must fill.\n"
    "4. Tie the rewrite back to the requirement it satisfies."
)


async def strengthen(requirement: dict, user_content: str,
                     company_profile: Optional[dict] = None,
                     solicitation_vocab: str = "") -> dict:
    """Rewrite the user's content to fully address one requirement — honestly."""
    profile = company_profile or {}
    certs = ", ".join(profile.get("certifications", []) or []) or "none stated"

    prompt = f"""Strengthen the offeror's content below so it fully and responsively addresses this requirement.

Return ONLY valid JSON:
{{
  "rewritten": "the strengthened section text (ready to paste)",
  "rationale": "1-2 sentences: how this now satisfies the requirement, in evaluator terms",
  "warnings": ["any fact the offeror must add/verify because it wasn't in their content — [] if none"],
  "invented_nothing": true
}}

REQUIREMENT ({requirement.get('id','')} — {requirement.get('category','')}):
{requirement.get('text','')}
SOURCE / EVALUATION LANGUAGE TO MIRROR: {requirement.get('source','')}
{('RELEVANT SOLICITATION VOCABULARY: ' + solicitation_vocab) if solicitation_vocab else ''}

OFFEROR'S CERTIFICATIONS (may reference if true): {certs}

OFFEROR'S EXISTING CONTENT (this is the ONLY factual basis you may use):
{_clip(user_content or '(the offeror provided no existing content for this requirement)', 12000)}

Remember: reorganize and sharpen their real content; never fabricate qualifications. If content is missing, say so in warnings instead of inventing."""
    resp = await _get_llm(4000).ainvoke([HumanMessage(content=f"{STRENGTHEN_SYSTEM}\n\n{prompt}")])
    data = _extract_json(_text_of(resp)) or {}
    return {
        "requirement_id": requirement.get("id", ""),
        "rewritten": data.get("rewritten", ""),
        "rationale": data.get("rationale", ""),
        "warnings": data.get("warnings", []) or [],
    }


# ── EXPORT: compliance matrix as a DOCX ──────────────────────────

_STATUS_LABEL = {"addressed": "Addressed", "partial": "Partial", "missing": "Missing"}
_STATUS_HEX = {"addressed": "1DB954", "partial": "F8C81C", "missing": "FF6432"}


def build_matrix_docx(title: str, requirements: List[dict], matrix: List[dict],
                      coverage_pct: int = 0, agency: str = "") -> io.BytesIO:
    """A submission-ready Requirements Compliance Matrix (Section L/M traceability)."""
    import docx
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    by_id = {str(m.get("id")): m for m in matrix}
    d = docx.Document()

    h = d.add_heading("Requirements Compliance Matrix", level=0)
    sub = d.add_paragraph()
    sub.add_run(title or "Solicitation").bold = True
    if agency:
        sub.add_run(f"  ·  {agency}")
    meta = d.add_paragraph()
    meta.add_run(f"Overall coverage: {coverage_pct}%   ·   {len(requirements)} requirements").italic = True

    table = d.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, label in enumerate(["ID", "Requirement", "Section", "Status", "Where addressed / gap"]):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(label)
        run.bold = True

    for r in requirements:
        m = by_id.get(r["id"], {})
        status = m.get("status", "missing")
        row = table.add_row().cells
        row[0].text = r.get("id", "")
        row[1].text = r.get("text", "")
        row[2].text = (r.get("section", "") or "").upper()
        # status cell — colored
        row[3].text = ""
        srun = row[3].paragraphs[0].add_run(_STATUS_LABEL.get(status, status.title()))
        srun.bold = True
        try:
            srun.font.color.rgb = RGBColor.from_string(_STATUS_HEX.get(status, "000000"))
        except Exception:
            pass
        note = m.get("note") or ""
        doc = m.get("doc")
        row[4].text = (f"[{doc}] " if doc else "") + note

    d.add_paragraph()
    foot = d.add_paragraph()
    foot.add_run("Generated by FinesseWins · finessewins.com").italic = True

    buf = io.BytesIO()
    d.save(buf)
    buf.seek(0)
    return buf
