"""
FinesseWins — DOCX export (python-docx)

Turns a stored proposal (its generated volumes) into a clean, submission-ready
Word document: cover page, per-volume sections with headings, page breaks
between volumes, and a footer with company / CAGE / UEI details.

The AI writes Markdown-ish prose (## headings, - bullets, **bold**). We do a
light parse so those render as real Word headings/lists instead of literal `##`.
"""
from __future__ import annotations

import io
import re
from datetime import datetime
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION

BRAND_MAGENTA = RGBColor(0xEC, 0x1C, 0x7B)
INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x66, 0x66, 0x66)

VOLUME_TITLES = {
    "technical": "Volume I — Technical Approach",
    "past_performance": "Volume II — Past Performance",
    "pricing": "Volume III — Price / Cost Proposal",
}


def build_proposal_docx(proposal: dict, profile: Optional[dict] = None) -> io.BytesIO:
    """Return a BytesIO containing the .docx for a proposal."""
    profile = profile or {}
    doc = Document()
    _base_styles(doc)
    _cover_page(doc, proposal, profile)

    volumes = proposal.get("volumes") or {}
    order = ["technical", "past_performance", "pricing"]
    for key in order:
        content = volumes.get(key)
        if not content:
            continue
        doc.add_page_break()
        _add_volume(doc, VOLUME_TITLES.get(key, key.replace("_", " ").title()), content)

    _footer(doc, profile)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ─────────────────────────────────────────────────────────────────────────────
def _base_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    for i, size in ((1, 16), (2, 13), (3, 11.5)):
        try:
            h = doc.styles[f"Heading {i}"]
            h.font.size = Pt(size)
            h.font.color.rgb = BRAND_MAGENTA if i == 1 else INK
            h.font.bold = True
        except KeyError:
            pass


def _cover_page(doc: Document, proposal: dict, profile: dict):
    for _ in range(3):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(proposal.get("title") or "Proposal")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = BRAND_MAGENTA

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(proposal.get("agency") or "")
    r.font.size = Pt(13)
    r.font.color.rgb = MUTED

    for _ in range(2):
        doc.add_paragraph()

    meta = [
        ("Solicitation No.", proposal.get("solicitation_number")),
        ("NAICS", proposal.get("naics_code")),
        ("Set-Aside", proposal.get("set_aside") or "None / Full & Open"),
        ("Response Deadline", _fmt_date(proposal.get("deadline"))),
    ]
    for label, value in meta:
        if not value:
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lr = p.add_run(f"{label}:  ")
        lr.font.bold = True
        lr.font.size = Pt(11)
        vr = p.add_run(str(value))
        vr.font.size = Pt(11)

    for _ in range(3):
        doc.add_paragraph()

    company = doc.add_paragraph()
    company.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = company.add_run(profile.get("name") or "Millennials Creatives LLC")
    cr.font.size = Pt(14)
    cr.font.bold = True

    certs = profile.get("certifications") or []
    if certs:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        crun = cp.add_run(" · ".join(certs))
        crun.font.size = Pt(10)
        crun.font.color.rgb = BRAND_MAGENTA

    ids = doc.add_paragraph()
    ids.alignment = WD_ALIGN_PARAGRAPH.CENTER
    id_bits = []
    if profile.get("uei"):
        id_bits.append(f"UEI {profile['uei']}")
    if profile.get("cage"):
        id_bits.append(f"CAGE {profile['cage']}")
    id_bits.append(f"Prepared {datetime.utcnow():%B %d, %Y}")
    ir = ids.add_run("  ·  ".join(id_bits))
    ir.font.size = Pt(9)
    ir.font.color.rgb = MUTED


def _add_volume(doc: Document, volume_title: str, content: str):
    doc.add_heading(volume_title, level=1)
    for block in _parse_markdownish(content):
        kind, text = block
        if kind == "h2":
            doc.add_heading(text, level=2)
        elif kind == "h3":
            doc.add_heading(text, level=3)
        elif kind == "bullet":
            doc.add_paragraph(text, style="List Bullet")
        elif kind == "number":
            doc.add_paragraph(text, style="List Number")
        elif kind == "blank":
            continue
        else:
            _rich_paragraph(doc, text)


def _rich_paragraph(doc: Document, text: str):
    """Paragraph with **bold** inline runs preserved."""
    p = doc.add_paragraph()
    for i, chunk in enumerate(re.split(r"(\*\*.+?\*\*)", text)):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            run = p.add_run(chunk[2:-2])
            run.font.bold = True
        else:
            p.add_run(chunk)


def _parse_markdownish(content: str):
    """Yield (kind, text) blocks from lightly-Markdown text."""
    for raw in content.splitlines():
        line = raw.rstrip()
        stripped = line.strip()
        if not stripped:
            yield ("blank", "")
        elif stripped.startswith("### "):
            yield ("h3", stripped[4:].strip())
        elif stripped.startswith("## "):
            yield ("h2", stripped[3:].strip())
        elif stripped.startswith("# "):
            yield ("h2", stripped[2:].strip())
        elif re.match(r"^[-*•]\s+", stripped):
            yield ("bullet", re.sub(r"^[-*•]\s+", "", stripped))
        elif re.match(r"^\d+[.)]\s+", stripped):
            yield ("number", re.sub(r"^\d+[.)]\s+", "", stripped))
        else:
            yield ("p", stripped)


def _footer(doc: Document, profile: dict):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bits = [profile.get("name") or "Millennials Creatives LLC"]
    if profile.get("cage"):
        bits.append(f"CAGE {profile['cage']}")
    if profile.get("uei"):
        bits.append(f"UEI {profile['uei']}")
    bits.append("Generated with FinesseWins")
    r = p.add_run("  ·  ".join(bits))
    r.font.size = Pt(8)
    r.font.color.rgb = MUTED


def _fmt_date(value) -> str:
    if not value:
        return ""
    try:
        dt = datetime.fromisoformat(str(value).replace("Z", "").replace("+00:00", ""))
        return dt.strftime("%B %d, %Y")
    except Exception:
        return str(value)


def build_capability_docx(content: str, profile: Optional[dict] = None) -> io.BytesIO:
    """Render a one-page capability statement (AI text) into a clean .docx."""
    profile = profile or {}
    doc = Document()
    _base_styles(doc)

    # Header: company name + certifications band
    title = doc.add_paragraph()
    run = title.add_run(profile.get("name") or "Capability Statement")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = BRAND_MAGENTA

    certs = profile.get("certifications") or []
    if certs:
        cp = doc.add_paragraph()
        cr = cp.add_run("  •  ".join(certs))
        cr.font.size = Pt(10)
        cr.font.color.rgb = BRAND_MAGENTA

    ids = []
    if profile.get("uei"):
        ids.append(f"UEI {profile['uei']}")
    if profile.get("cage"):
        ids.append(f"CAGE {profile['cage']}")
    if profile.get("naics_codes"):
        ids.append("NAICS " + ", ".join(profile["naics_codes"][:6]))
    if ids:
        ip = doc.add_paragraph()
        ir = ip.add_run("  ·  ".join(ids))
        ir.font.size = Pt(9)
        ir.font.color.rgb = MUTED

    # horizontal rule-ish spacer
    doc.add_paragraph()

    for block in _parse_markdownish(content or ""):
        kind, text = block
        if kind == "h2":
            doc.add_heading(text, level=2)
        elif kind == "h3":
            doc.add_heading(text, level=3)
        elif kind == "bullet":
            doc.add_paragraph(text, style="List Bullet")
        elif kind == "number":
            doc.add_paragraph(text, style="List Number")
        elif kind == "blank":
            continue
        else:
            _rich_paragraph(doc, text)

    _footer(doc, profile)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def safe_filename(proposal: dict) -> str:
    base = (proposal.get("solicitation_number") or proposal.get("title") or "proposal")
    base = re.sub(r"[^A-Za-z0-9._-]+", "_", str(base)).strip("_")
    return f"{base or 'proposal'}.docx"
