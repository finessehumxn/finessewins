"""
Document parsing for the RFP shredder.

Turns an uploaded PDF / DOCX / TXT into (a) clean full text and (b) a list of
titled sections, so the shredder can point every extracted requirement back to
where it came from, and the gap analysis can score the user's material
section-by-section.

Pure parsing — no AI, no network — so it's fast and unit-testable offline.
"""
from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from typing import List


@dataclass
class Section:
    heading: str
    text: str


@dataclass
class ParsedDoc:
    filename: str
    kind: str                       # "pdf" | "docx" | "txt"
    text: str
    sections: List[Section] = field(default_factory=list)
    pages: int = 0

    @property
    def word_count(self) -> int:
        return len(self.text.split())


# Lines that read like a heading in a government solicitation:
#   "SECTION L - INSTRUCTIONS", "L.1 General", "3.2 Scope of Work",
#   "ATTACHMENT 2B", "EVALUATION CRITERIA", "M.2.1 ..."
_HEADING_RE = re.compile(
    r"^(?:"
    r"(?:SECTION\s+[A-Z0-9]|ATTACHMENT\s+\w|EXHIBIT\s+\w|APPENDIX\s+\w)"      # SECTION L / ATTACHMENT 2B
    r"|[A-Z]\.\d+(?:\.\d+)*\b"                                                  # L.1  M.2.1
    r"|\d+(?:\.\d+){0,3}\s+[A-Z]"                                               # 3.2 Scope
    r"|[A-Z][A-Z0-9 ,/&'\-]{6,60}"                                             # ALL-CAPS TITLE
    r")"
)


def _looks_like_heading(line: str) -> bool:
    s = line.strip()
    if not (3 <= len(s) <= 90):
        return False
    if s.endswith((".", ",", ";", ":")) and s.isupper() is False and not _HEADING_RE.match(s):
        return False
    # Mostly-uppercase short line, or a numbered/section prefix.
    if _HEADING_RE.match(s):
        # Avoid treating a normal sentence in Title Case as a heading:
        letters = [c for c in s if c.isalpha()]
        upper_ratio = sum(c.isupper() for c in letters) / max(1, len(letters))
        return upper_ratio > 0.55 or bool(re.match(r"^(?:SECTION|ATTACHMENT|EXHIBIT|APPENDIX)\b", s, re.I)) \
            or bool(re.match(r"^[A-Z]\.\d", s)) or bool(re.match(r"^\d+(?:\.\d+)*\s+[A-Z]", s))
    return False


def _split_sections(text: str) -> List[Section]:
    """Group lines under the nearest preceding heading."""
    sections: List[Section] = []
    cur_head = "Document"
    cur_lines: List[str] = []
    for raw in text.splitlines():
        line = raw.rstrip()
        if _looks_like_heading(line):
            if cur_lines and any(l.strip() for l in cur_lines):
                sections.append(Section(cur_head, "\n".join(cur_lines).strip()))
            cur_head = line.strip()
            cur_lines = []
        else:
            cur_lines.append(line)
    if cur_lines and any(l.strip() for l in cur_lines):
        sections.append(Section(cur_head, "\n".join(cur_lines).strip()))
    # If nothing split (flat doc), keep it as one section.
    if not sections:
        sections = [Section("Document", text.strip())]
    return sections


def _clean(text: str) -> str:
    # Collapse runs of blank lines and stray form-feeds; keep single newlines.
    text = text.replace("\x0c", "\n")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _parse_pdf(data: bytes, filename: str) -> ParsedDoc:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    parts = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            parts.append("")
    text = _clean("\n".join(parts))
    return ParsedDoc(filename=filename, kind="pdf", text=text,
                     sections=_split_sections(text), pages=len(reader.pages))


def _parse_docx(data: bytes, filename: str) -> ParsedDoc:
    import docx
    d = docx.Document(io.BytesIO(data))
    lines: List[str] = []
    sections: List[Section] = []
    cur_head = "Document"
    cur_lines: List[str] = []

    def flush():
        nonlocal cur_lines, cur_head
        if cur_lines and any(l.strip() for l in cur_lines):
            sections.append(Section(cur_head, "\n".join(cur_lines).strip()))
        cur_lines = []

    for p in d.paragraphs:
        txt = (p.text or "").rstrip()
        style = (p.style.name or "") if p.style else ""
        is_head = style.lower().startswith("heading") or style.lower() == "title"
        lines.append(txt)
        if is_head and txt.strip():
            flush()
            cur_head = txt.strip()
        elif _looks_like_heading(txt):
            flush()
            cur_head = txt.strip()
        else:
            cur_lines.append(txt)
    flush()

    # Pull table cell text too (gov docs love tables) into the full text.
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append(" | ".join(cells))

    text = _clean("\n".join(lines))
    if not sections:
        sections = _split_sections(text)
    return ParsedDoc(filename=filename, kind="docx", text=text, sections=sections)


def parse(data: bytes, filename: str) -> ParsedDoc:
    """Parse an uploaded document by extension. Raises ValueError on unsupported."""
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return _parse_pdf(data, filename)
    if name.endswith(".docx"):
        return _parse_docx(data, filename)
    if name.endswith((".txt", ".md")):
        text = _clean(data.decode("utf-8", errors="replace"))
        return ParsedDoc(filename=filename, kind="txt", text=text, sections=_split_sections(text))
    raise ValueError(f"Unsupported file type: {filename}. Upload a PDF, DOCX, or TXT.")
